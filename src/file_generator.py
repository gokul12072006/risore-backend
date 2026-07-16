import json
import requests
from datetime import datetime
from docx import Document
from markdown_pdf import MarkdownPdf, Section


def generate_pdf(title: str, content: str, filepath: str) -> bool:
    try:
        pdf = MarkdownPdf(toc_level=2)
        full_content = f"# {title}\n\n{content}"
        pdf.add_section(Section(full_content))
        pdf.save(filepath)
        return True
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return False


def generate_word(title: str, content: str, filepath: str) -> bool:
    try:
        doc = Document()
        doc.add_heading(title, level=1)

        # Add content, handling line breaks
        for paragraph in content.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        doc.save(filepath)
        return True
    except Exception as e:
        print(f"Error generating Word document: {e}")
        return False


def generate_weather_svg(city: str, filepath: str) -> bool:
    try:
        res = requests.get(f"https://wttr.in/{city}?format=j1", timeout=10)
        if res.status_code != 200:
            return False
        data = res.json()
        
        curr = data['current_condition'][0]
        temp_c = curr['temp_C']
        desc = curr['weatherDesc'][0]['value']
        feels = curr['FeelsLikeC']
        wind = curr['windspeedKmph']
        humidity = curr['humidity']
        
        forecasts = data['weather'][:3]
        
        svg = f'''<svg width="450" height="280" viewBox="0 0 450 280" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <linearGradient id="bg" x1="0%" y1="0%" x2="100%" y2="100%">
      <stop offset="0%" stop-color="#1e1e24" />
      <stop offset="100%" stop-color="#131318" />
    </linearGradient>
    <filter id="shadow">
      <feDropShadow dx="0" dy="4" stdDeviation="6" flood-opacity="0.3" />
    </filter>
  </defs>
  <rect width="450" height="280" rx="20" fill="url(#bg)" filter="url(#shadow)" stroke="#33333b" stroke-width="1.5" />
  
  <text x="25" y="45" font-family="system-ui, -apple-system, sans-serif" font-size="22" font-weight="bold" fill="#ffffff">{city.upper()}</text>
  <text x="25" y="65" font-family="system-ui, -apple-system, sans-serif" font-size="14" fill="#a0a0a5">{desc} • Feels like {feels}°C</text>
  
  <text x="25" y="140" font-family="system-ui, -apple-system, sans-serif" font-size="64" font-weight="bold" fill="#ffffff">{temp_c}°</text>
  <text x="135" y="115" font-family="system-ui, -apple-system, sans-serif" font-size="20" font-weight="bold" fill="#4da5fc">C</text>
  
  <text x="25" y="180" font-family="system-ui, -apple-system, sans-serif" font-size="14" fill="#a0a0a5">Humidity: {humidity}%  |  Wind: {wind} km/h</text>
  
  <line x1="25" y1="200" x2="425" y2="200" stroke="#33333b" stroke-width="1" />
  
  <!-- Forecast -->
  '''
        x_offset = 35
        for day in forecasts:
            date_str = day['date']
            try:
                dt = datetime.strptime(date_str, "%Y-%m-%d")
                day_name = dt.strftime("%A") if dt.date() != datetime.now().date() else "Today"
            except:
                day_name = "Day"
            
            max_t = day['maxtempC']
            min_t = day['mintempC']
            
            svg += f'''
  <text x="{x_offset}" y="230" font-family="system-ui, -apple-system, sans-serif" font-size="14" font-weight="bold" fill="#ffffff">{day_name}</text>
  <text x="{x_offset}" y="255" font-family="system-ui, -apple-system, sans-serif" font-size="13" fill="#a0a0a5">{max_t}° / {min_t}°</text>
'''
            x_offset += 140

        svg += "</svg>"

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(svg)
        return True
    except Exception as e:
        print(f"Weather SVG Error: {e}")
        return False
